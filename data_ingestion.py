import pandas as pd
import json
import streamlit as st
import io
import PyPDF2
from docx import Document
uploaded_file = st.file_uploader(
    "Upload File",
    type=["csv", "xlsx", "xls", "pdf", "docx", "json", "txt"]
)

def handle_file(uploaded_file):
    file_type = uploaded_file.name.split(".")[-1].lower()

    # ================= CSV =================
    if file_type == "csv":
        df = pd.read_csv(uploaded_file)
        return df

    # ================= EXCEL =================
    elif file_type in ["xlsx", "xls"]:
        df = pd.read_excel(uploaded_file)
        return df

    # ================= JSON =================
    elif file_type == "json":
        data = json.load(uploaded_file)
        if isinstance(data, list):
            return pd.DataFrame(data)
        else:
            return pd.DataFrame([data])

    # ================= TXT =================
    elif file_type == "txt":
        text = uploaded_file.read().decode("utf-8")
        return text

    # ================= PDF =================
    elif file_type == "pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

    # ================= WORD =================
    elif file_type == "docx":
        doc = Document(uploaded_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text

    else:
        return None

# def hand(uploaded_file):
#     """
#     Universal file handler.
#     Extracts content and displays it inside a scrollable window.
#     """

#     if uploaded_file is None:
#         return

#     file_type = uploaded_file.name.split(".")[-1].lower()

#     try:
#         content = ""

#         # ================= CSV =================
#         if file_type == "csv":
#             df = pd.read_csv(uploaded_file)
#             df.fillna("", inplace=True)
#             content = df.to_csv(index=False)

#         # ================= EXCEL =================
#         elif file_type in ["xlsx", "xls"]:
#             df = pd.read_excel(uploaded_file)
#             df.fillna("", inplace=True)
#             content = df.to_csv(index=False)

#         # ================= JSON =================
#         elif file_type == "json":
#             data = json.load(uploaded_file)
#             content = json.dumps(data, indent=2)

#         # ================= TXT =================
#         elif file_type == "txt":
#             content = uploaded_file.read().decode("utf-8", errors="ignore")

#         # ================= PDF =================
#         elif file_type == "pdf":
#             pdf_reader = PyPDF2.PdfReader(uploaded_file)
#             for page in pdf_reader.pages:
#                 extracted = page.extract_text()
#                 if extracted:
#                     content += extracted + "\n"

#         # ================= WORD =================
#         elif file_type == "docx":
#             doc = Document(uploaded_file)
#             content = "\n".join(
#                 [para.text for para in doc.paragraphs if para.text.strip()]
#             )

#         else:
#             st.error("Unsupported file type.")
#             return
#     except Exception as e:
#         return f"Error processing file: {str(e)}"

def hand(uploaded_file):
    if uploaded_file is None:
        return ""

    # CRITICAL: Reset the file pointer so it can be read multiple times
    uploaded_file.seek(0)
    
    file_type = uploaded_file.name.split(".")[-1].lower()

    try:
        content = ""

        # CSV / EXCEL
        if file_type in ["csv", "xlsx", "xls"]:
            import pandas as pd
            if file_type == "csv":
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            df.fillna("", inplace=True)
            
            content = df#to_csv(index=False)

        # JSON
        elif file_type == "json":
            import json
            data = json.load(uploaded_file)
            content = json.dumps(data, indent=2)

        # TXT
        elif file_type == "txt":
            content = uploaded_file.read().decode("utf-8", errors="ignore")

        # PDF
        elif file_type == "pdf":
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text_parts = []
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                text_parts.append(extracted if extracted else "[Unreadable Page/Image]")
            content = "\n".join(text_parts)

        # WORD
        elif file_type == "docx":
            from docx import Document
            doc = Document(uploaded_file)
            content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

        else:
            return "Unsupported file type."

        return content,file_type

    except Exception as e:
        return f"Error processing {uploaded_file.name}: {str(e)}"

def handled_file(uploaded_file):
    """
    Universal file handler.
    Returns CLEAN TEXT for AI processing.
    """

    if uploaded_file is None:
        return None

    file_type = uploaded_file.name.split(".")[-1].lower()

    try:

        # ================= CSV =================
        if file_type == "csv":
            df = pd.read_csv(uploaded_file)
            df.fillna("", inplace=True)
            return df.to_csv(index=False)

        # ================= EXCEL =================
        elif file_type in ["xlsx", "xls"]:
            df = pd.read_excel(uploaded_file)
            df.fillna("", inplace=True)
            return df.to_csv(index=False)

        # ================= JSON =================
        elif file_type == "json":
            data = json.load(uploaded_file)
            return json.dumps(data, indent=2)

        # ================= TXT =================
        elif file_type == "txt":
            return uploaded_file.read().decode("utf-8", errors="ignore")

        # ================= PDF =================
        elif file_type == "pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text

        # ================= WORD =================
        elif file_type == "docx":
            doc = Document(uploaded_file)
            text = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip() != ""]
            )
            return text

        else:
            return "Unsupported file type."

    except Exception as e:
        return f"Error processing file: {str(e)}"
# def scrollable_window(content,file_type):
#     st.markdown("### File Preview")
#     if file_type in ["xlsx","xls","csv"]:
#         st.markdown(
#             f"""
#             <div style="
#                 height:500px;
#                 overflow:auto;
#                 border:1px solid #f5c6cb;
#                 padding:15px;
#                 border-radius:10px;
#                 color: #dc3545;
#                 background-color:#f8f9fa;
#                 white-space: pre-wrap;
#             ">
#             {st.dataframe(content)}
#             </div>
#             """,
#             unsafe_allow_html=True
#         )
#     else:
#         st.markdown(
#             f"""
#             <div style="
#                 height:500px;
#                 overflow:auto;
#                 border:1px solid #f5c6cb;
#                 padding:15px;
#                 border-radius:10px;
#                 color: #dc3545;
#                 background-color:#f8f9fa;
#                 white-space: pre-wrap;
#             ">
#             {content}
#             </div>
#             """,
#             unsafe_allow_html=True
#         )
def scrollable_window(content, file_type):
    st.markdown("### File Preview")
    
    # Check if we are dealing with a table
    if file_type in ["xlsx", "xls", "csv"]:
        # st.dataframe provides its own scrollbars and interface
        # No need for the <div> wrapper here
        st.dataframe(content, use_container_width=True, height=500)
        
    else:
        # For text-based files (PDF, TXT, DOCX), use your custom scrollable div
        st.markdown(
            f"""
            <div style="
                height:500px;
                overflow:auto;
                border:1px solid #f5c6cb;
                padding:15px;
                border-radius:10px;
                color: #dc3545;
                background-color:#f8f9fa;
                white-space: pre-wrap;
            ">
            {content}
            </div>
            """,
            unsafe_allow_html=True
        )


    return content
# def scrollable_window(content):
#     st.markdown("### File Preview")

#     # Define the CSS separately to keep it clean
#     html_style = """
#         <div style="
#             height:500px;
#             overflow:auto;
#             border:1px solid #f5c6cb;
#             padding:15px;
#             border-radius:10px;
#             color: #dc3545;
#             background-color:#f8f9fa;
#             white-space: pre-wrap;
#         ">
#     """
    
#     # Combine them safely without using an f-string on the content
#     # full_html = html_style + str(content) + "</div>"

#     st.markdown(content, unsafe_allow_html=True)
#     return content


def handling_file(uploaded_file):
    """
    Universal file handler.
    Extracts content and displays it inside a scrollable window.
    """

    if uploaded_file is None:
        return

    file_type = uploaded_file.name.split(".")[-1].lower()

    try:
        content = ""

        # ================= CSV =================
        if file_type == "csv":
            df = pd.read_csv(uploaded_file)
            df.fillna("", inplace=True)
            content = df.to_csv(index=False)

        # ================= EXCEL =================
        elif file_type in ["xlsx", "xls"]:
            df = pd.read_excel(uploaded_file)
            df.fillna("", inplace=True)
            content = df.to_csv(index=False)

        # ================= JSON =================
        elif file_type == "json":
            data = json.load(uploaded_file)
            content = json.dumps(data, indent=2)

        # ================= TXT =================
        elif file_type == "txt":
            content = uploaded_file.read().decode("utf-8", errors="ignore")

        # ================= PDF =================
        elif file_type == "pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                extracted = page.extract_text()
                if extracted:
                    content += extracted + "\n"

        # ================= WORD =================
        elif file_type == "docx":
            doc = Document(uploaded_file)
            content = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()]
            )

        else:
            st.error("Unsupported file type.")
            return

        # ================= SCROLLABLE WINDOW =================
        st.markdown("### File Preview")

        st.markdown(
            f"""
            <div style="
                height:500px;
                overflow:auto;
                border:1px solid #ccc;
                padding:15px;
                border-radius:10px;
                background-color:#f8f9fa;
                white-space: pre-wrap;
            ">
            {content}
            </div>
            """,
            unsafe_allow_html=True
        )

        return content

    except Exception as e:
        st.error(f"Error processing file: {str(e)}")



